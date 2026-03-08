# ==================== app.py ====================
import streamlit as st
import docx
import re
from io import BytesIO
from datetime import datetime
import dashscope
from dashscope import Generation

# 页面配置
st.set_page_config(page_title="《下一站》剪辑稿生成器", layout="wide")

# 标题
st.title("🎬 《下一站》剪辑稿智能生成系统")
st.markdown("---")

# 初始化会话状态
if 'field_notes_lines' not in st.session_state:
    st.session_state.field_notes_lines = []
if 'selected_clips' not in st.session_state:
    st.session_state.selected_clips = []
if 'segments' not in st.session_state:
    st.session_state.segments = []
if 'api_key' not in st.session_state:
    st.session_state.api_key = ""

# ==================== 侧边栏 ====================
with st.sidebar:
    st.header("📖 使用说明")
    st.write("1. 上传场记稿文件")
    st.write("2. **复制一小句内容**→粘贴到标记区→点击标记")
    st.write("3. 添加分段，将素材**按序号分配**到各分段")
    st.write("4. 可**拖动调整序号**改变语序")
    st.write("5. 生成并下载两个文档")
    
    st.header("📋 格式规范")
    st.write("- OS旁白：`OS：内容//`")
    st.write("- 实况对话：`内容//` (自动添加)")
    st.write("- 同一段内不换行，连续排列")
    st.write("- 分段时换行，用OS衔接")
    
    st.header("⚙️ AI设置")
    api_key_input = st.text_input("通义千问API Key", type="password", value=st.session_state.api_key)
    if api_key_input:
        st.session_state.api_key = api_key_input
        dashscope.api_key = api_key_input

# ==================== 第1步：上传场记稿 ====================
st.header("📄 第1步：上传场记稿")
uploaded_file = st.file_uploader("上传场记稿(.docx)", type=["docx"])

if uploaded_file:
    # 读取文档
    doc = docx.Document(uploaded_file)
    full_text = []
    for para in doc.paragraphs:
        if para.text.strip():
            full_text.append(para.text)
    
    field_notes_text = "\n".join(full_text)
    
    # 按行分割（每行一个可标记单元）
    lines = field_notes_text.split('\n')
    st.session_state.field_notes_lines = lines
    
    with st.expander("📋 场记稿预览"):
        st.text_area("原始内容", field_notes_text, height=200)

# ==================== 第2步：精细高光标记 ====================
st.header("✏️ 第2步：高光标记（复制一小句内容标记）")

if st.session_state.field_notes_lines:
    st.info("💡 **操作提示**：在场记稿预览中**复制一小句内容**→粘贴到下方输入框→点击标记")
    
    col1, col2 = st.columns([2, 1])
    
    with col1:
        # 复制粘贴区域 - 精细化标记
        selected_text = st.text_area(
            "📋 粘贴要标记的内容（一小句）", 
            height=80, 
            key="selected_text",
            placeholder="例如：你要喝点什么吗？"
        )
    
    with col2:
        # 元数据输入
        timecode = st.text_input("时间码", placeholder="如：3:59")
        speaker = st.text_input("讲话人", placeholder="如：讲话人1")
        location = st.text_input("地点/场景", placeholder="如：咖啡巴士")
    
    # 标记按钮
    col_btn1, col_btn2, col_btn3 = st.columns([2, 1, 1])
    
    with col_btn1:
        if st.button("✅ 标记选中内容", use_container_width=True, type="primary"):
            if selected_text and selected_text.strip():
                # 检查是否重复
                is_duplicate = False
                for clip in st.session_state.selected_clips:
                    if clip['content'] == selected_text.strip():
                        is_duplicate = True
                        break
                
                if is_duplicate:
                    st.warning("⚠️ 该内容已标记过，请勿重复添加")
                else:
                    clip = {
                        "id": len(st.session_state.selected_clips) + 1,
                        "content": selected_text.strip(),
                        "timecode": timecode,
                        "speaker": speaker,
                        "location": location,
                        "type": "实况",
                        "notes": "",
                        "original_line_index": -1  # 记录在原稿中的位置
                    }
                    
                    # 尝试匹配原稿中的行
                    for i, line in enumerate(st.session_state.field_notes_lines):
                        if selected_text.strip() in line:
                            clip['original_line_index'] = i
                            break
                    
                    st.session_state.selected_clips.append(clip)
                    st.success(f"✅ 已标记！当前共 {len(st.session_state.selected_clips)} 段素材")
    
    with col_btn2:
        if st.button("🗑️ 清空所有", use_container_width=True):
            st.session_state.selected_clips = []
            st.success("已清空")
    
    with col_btn3:
        if st.button("📋 批量导入", use_container_width=True):
            st.info("批量导入功能开发中...")

# 显示已标记素材库
if st.session_state.selected_clips:
    st.subheader(f"📦 已标记素材库 ({len(st.session_state.selected_clips)}段)")
    
    # 搜索过滤
    search_term = st.text_input("🔍 搜索素材", placeholder="输入关键词搜索...")
    
    # 显示素材列表（带序号）
    cols = st.columns([1, 6, 2, 1])
    with cols[0]:
        st.write("**序号**")
    with cols[1]:
        st.write("**内容**")
    with cols[2]:
        st.write("**时间码**")
    with cols[3]:
        st.write("**操作**")
    
    for i, clip in enumerate(st.session_state.selected_clips):
        # 过滤搜索
        if search_term and search_term not in clip['content']:
            continue
        
        row_cols = st.columns([1, 6, 2, 1])
        
        with row_cols[0]:
            st.write(f"`{i+1}`")
        
        with row_cols[1]:
            # 可编辑内容
            new_content = st.text_area(
                "内容", 
                clip['content'], 
                height=60, 
                key=f"edit_{i}",
                label_visibility="collapsed"
            )
            if new_content != clip['content']:
                clip['content'] = new_content
            
            # 备注
            notes = st.text_input("备注", clip['notes'], key=f"notes_{i}", placeholder="如：调整语序")
            clip['notes'] = notes
        
        with row_cols[2]:
            st.write(clip['timecode'] or "-")
        
        with row_cols[3]:
            # 上移
            if st.button("⬆️", key=f"up_{i}", help="上移"):
                if i > 0:
                    st.session_state.selected_clips[i], st.session_state.selected_clips[i-1] = \
                    st.session_state.selected_clips[i-1], st.session_state.selected_clips[i]
                    st.rerun()
            
            # 删除
            if st.button("❌", key=f"del_{i}", help="删除"):
                st.session_state.selected_clips.pop(i)
                st.rerun()

# ==================== 第3步：分段管理（带序号） ====================
st.header("📑 第3步：分段管理（带序号排列）")

if st.session_state.selected_clips:
    # 手动添加分段
    col_seg1, col_seg2 = st.columns([3, 1])
    with col_seg1:
        new_segment_name = st.text_input("新分段名称（地点/主题）", placeholder="如：咖啡巴士引入/红枫林观景")
    with col_seg2:
        if st.button("➕ 添加分段"):
            if new_segment_name:
                st.session_state.segments.append({
                    "name": new_segment_name,
                    "clip_indices": [],  # 存储素材序号列表
                    "os_suggestion": "",
                    "os_custom": ""
                })
                st.success(f"已添加分段：{new_segment_name}")
    
    # 显示分段
    if st.session_state.segments:
        st.subheader("当前分段")
        for i, seg in enumerate(st.session_state.segments):
            with st.expander(f"分段{i+1}: {seg['name']}", expanded=True):
                # 分段说明
                st.info(f"💡 从上方素材库中选择素材，**按序号添加**到此分段。序号决定最终剪辑顺序。")
                
                # 可用素材选择（显示序号）
                st.write("**可用素材：**")
                available_cols = st.columns(3)
                for j, clip in enumerate(st.session_state.selected_clips):
                    col_idx = j % 3
                    with available_cols[col_idx]:
                        if st.checkbox(
                            f"`{j+1}` {clip['content'][:30]}...", 
                            key=f"seg_{i}_avail_{j}",
                            value=j in seg['clip_indices']
                        ):
                            if j not in seg['clip_indices']:
                                seg['clip_indices'].append(j)
                        else:
                            if j in seg['clip_indices']:
                                seg['clip_indices'].remove(j)
                
                # 已选素材（带序号，可调整顺序）
                if seg['clip_indices']:
                    st.write("**✅ 已选素材（按序号排列）：**")
                    
                    # 显示已选素材列表，可调整顺序
                    for idx_pos, clip_idx in enumerate(seg['clip_indices']):
                        clip = st.session_state.selected_clips[clip_idx]
                        
                        seg_row_cols = st.columns([1, 5, 2, 2])
                        
                        with seg_row_cols[0]:
                            st.write(f"`{idx_pos+1}`")
                        
                        with seg_row_cols[1]:
                            st.write(f"{clip['content'][:50]}...")
                        
                        with seg_row_cols[2]:
                            # 上移
                            if st.button("⬆️", key=f"seg_{i}_up_{idx_pos}"):
                                if idx_pos > 0:
                                    seg['clip_indices'][idx_pos], seg['clip_indices'][idx_pos-1] = \
                                    seg['clip_indices'][idx_pos-1], seg['clip_indices'][idx_pos]
                                    st.rerun()
                            
                            # 下移
                            if st.button("⬇️", key=f"seg_{i}_down_{idx_pos}"):
                                if idx_pos < len(seg['clip_indices']) - 1:
                                    seg['clip_indices'][idx_pos], seg['clip_indices'][idx_pos+1] = \
                                    seg['clip_indices'][idx_pos+1], seg['clip_indices'][idx_pos]
                                    st.rerun()
                        
                        with seg_row_cols[3]:
                            # 移除
                            if st.button("❌", key=f"seg_{i}_remove_{idx_pos}"):
                                seg['clip_indices'].remove(clip_idx)
                                st.rerun()
                
                # OS建议
                st.write("**🎙️ OS旁白（探访人视角）：**")
                use_ai = st.checkbox("使用AI生成OS建议", key=f"ai_{i}")
                
                if use_ai and st.session_state.api_key:
                    if st.button("🤖 生成OS建议", key=f"gen_os_{i}"):
                        # 获取此分段的内容
                        seg_contents = [st.session_state.selected_clips[idx]['content'] 
                                       for idx in seg['clip_indices'][:5]]
                        context = "\n".join(seg_contents)
                        
                        # 调用通义千问API
                        prompt = f"""你是《下一站》电视节目编导，请根据以下采访内容，写一句探访人视角的OS旁白。
要求：
1. 以第一人称"我"的角度
2. 可以是现场情感表达或补充说明
3. 简洁有力，30字以内
4. 用于衔接上下两段内容

采访内容：
{context}

请直接输出OS内容，不要其他说明："""
                        
                        try:
                            response = Generation.call(
                                model='qwen-turbo',
                                prompt=prompt
                            )
                            if response.status_code == 200:
                                seg['os_suggestion'] = response.output.text.strip()
                                st.success("OS建议已生成！")
                        except Exception as e:
                            st.error(f"AI生成失败：{str(e)}")
                    elif seg['os_suggestion']:
                        st.info(f"AI建议：{seg['os_suggestion']}")
                elif use_ai and not st.session_state.api_key:
                    st.warning("请先在侧边栏输入通义千问API Key")
                
                # 手动编辑OS
                os_custom = st.text_area(
                    "OS旁白（可编辑）", 
                    seg['os_custom'] if seg['os_custom'] else seg['os_suggestion'], 
                    height=60,
                    placeholder="如：边喝咖啡，边欣赏沿路美景...", 
                    key=f"os_custom_{i}"
                )
                seg['os_custom'] = os_custom
                
                # 删除分段
                if st.button(f"🗑️ 删除此分段", key=f"del_seg_{i}"):
                    st.session_state.segments.pop(i)
                    st.rerun()
    else:
        st.info("请先添加分段，然后将素材分配到各分段")

# ==================== 第4步：预览和导出 ====================
st.header("💾 第4步：预览和导出")

if st.button("📝 生成剪辑稿预览"):
    if st.session_state.segments:
        # 生成剪辑稿内容
        edit_script = ""
        
        for seg_idx, seg in enumerate(st.session_state.segments):
            # 添加分段标题
            edit_script += f"【分段{seg_idx+1}: {seg['name']}】\n"
            
            # 添加OS（优先使用手动编辑的，其次使用AI建议）
            os_text = seg['os_custom'] if seg['os_custom'] else seg['os_suggestion']
            if os_text:
                edit_script += f"OS：{os_text}//\n"
            
            # 添加实况内容（按序号排列，同一段内不换行，连续排列）
            segment_content = ""
            for clip_idx in seg['clip_indices']:
                if clip_idx < len(st.session_state.selected_clips):
                    clip = st.session_state.selected_clips[clip_idx]
                    
                    # 添加备注
                    if clip['notes']:
                        edit_script += f"【{clip['notes']}】\n"
                    
                    # 添加内容（自动添加//，不换行）
                    segment_content += f"{clip['content']}// "
            
            if segment_content:
                edit_script += f"{segment_content}\n"
            
            edit_script += "\n" + "="*50 + "\n\n"
        
        st.session_state.edit_script_preview = edit_script
        
        # 显示预览
        st.subheader("剪辑稿预览")
        st.text_area("预览内容", edit_script, height=400)
    else:
        st.warning("请先添加分段并分配内容")

# 导出按钮
col_export1, col_export2 = st.columns(2)

with col_export1:
    if st.button("📥 下载剪辑稿(.docx)", use_container_width=True, type="primary"):
        if 'edit_script_preview' in st.session_state:
            # 创建Word文档
            doc = docx.Document()
            doc.add_heading('《下一站》剪辑稿', 0)
            doc.add_paragraph(f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M')}")
            doc.add_paragraph(st.session_state.edit_script_preview)
            
            # 保存到内存
            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            
            # 提供下载
            st.download_button(
                label="点击下载剪辑稿",
                data=buffer,
                file_name=f"剪辑稿_{datetime.now().strftime('%Y%m%d')}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
        else:
            st.warning("请先生成预览")

with col_export2:
    if st.button("📥 下载标记场记稿(.docx)", use_container_width=True):
        if uploaded_file and st.session_state.selected_clips:
            # 创建标记后的场记稿
            doc = docx.Document()
            doc.add_heading('《下一站》场记稿（已标记）', 0)
            doc.add_paragraph(f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M')}")
            doc.add_paragraph(f"已标记内容数量：{len(st.session_state.selected_clips)}\n")
            
            # 添加原始内容，高亮标记部分
            doc.add_heading('原始场记稿（✅=已标记）', level=1)
            for i, line in enumerate(st.session_state.field_notes_lines):
                if line.strip():
                    # 检查该行是否包含已标记内容
                    is_marked = False
                    marked_content = ""
                    for clip in st.session_state.selected_clips:
                        if clip['content'] in line:
                            is_marked = True
                            marked_content = clip['content']
                            break
                    
                    if is_marked:
                        p = doc.add_paragraph()
                        # 高亮显示标记内容
                        parts = line.split(marked_content)
                        for j, part in enumerate(parts):
                            if part:
                                p.add_run(part)
                            if j < len(parts) - 1:
                                runner = p.add_run(marked_content)
                                runner.font.highlight_color = docx.enum.text.WD_COLOR_INDEX.YELLOW
                    else:
                        doc.add_paragraph(line)
            
            # 添加标记内容列表
            doc.add_heading('已标记高光内容清单', level=1)
            for i, clip in enumerate(st.session_state.selected_clips):
                p = doc.add_paragraph()
                p.add_run(f"【标记{i+1}】").bold = True
                p.add_run(f" 时间码：{clip['timecode'] or '-'}  ")
                p.add_run(f"讲话人：{clip['speaker'] or '-'}  ")
                p.add_run(f"地点：{clip['location'] or '-'}\n")
                p.add_run(clip['content'])
            
            # 保存到内存
            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            
            # 提供下载
            st.download_button(
                label="点击下载标记场记稿",
                data=buffer,
                file_name=f"场记稿_已标记_{datetime.now().strftime('%Y%m%d')}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
        else:
            st.warning("请先上传场记稿并标记内容")

# 页脚
st.markdown("---")
st.caption("《下一站》节目组内部工具 | 版本 3.0 | 支持精细标记+序号排列")
# ==================== 代码结束 ====================
